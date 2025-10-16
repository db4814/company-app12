# -*- coding: utf-8 -*-
# 依赖测试代码 - 放在所有import之前
print("=== 依赖包测试 ===")
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    print("✅ python-docx 安装成功！")
except ImportError as e:
    print(f"❌ python-docx 安装失败: {e}")

try:
    from reportlab.lib.pagesizes import A4
    print("✅ reportlab 安装成功！")
except ImportError as e:
    print(f"❌ reportlab 安装失败: {e}")
print("==================")

# 原有的import继续...
from flask import Flask, render_template, request, jsonify, send_file
import sqlite3
import os
import pandas as pd
import io
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import base64

# 创建app实例
app = Flask(__name__)
app.secret_key = 'enterprise-system-v2-2024'

# ==================== 模板函数 ====================
@app.context_processor
def utility_processor():
    def get_data_type_name(data_type):
        names = {
            'output': '产值',
            'capacity': '产能', 
            'tax': '税收',
            'investment': '固定投资',
            'added_value': '增加值'
        }
        return names.get(data_type, '数据')

    def get_data_type_unit(data_type):
        units = {
            'output': '万元',
            'capacity': '',
            'tax': '万元', 
            'investment': '万元',
            'added_value': '万元'
        }
        return units.get(data_type, '')

    return dict(get_data_type_name=get_data_type_name, get_data_type_unit=get_data_type_unit)

# ==================== 设置响应头 ====================
@app.after_request
def after_request(response):
    response.headers['Content-Type'] = 'text/html; charset=utf-8'
    return response

def get_db_connection():
    conn = sqlite3.connect('enterprise.db')
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    
    # 企业基本信息表
    conn.execute('''
        CREATE TABLE IF NOT EXISTS companies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            legal_person TEXT,
            main_products TEXT,
            product_model TEXT,
            party_secretary TEXT,
            total_investment REAL DEFAULT 0,
            employee_count INTEGER DEFAULT 0,
            register_date TEXT,
            completion_date TEXT,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # 联系人表（法人、书记、日常联系人）
    conn.execute('''
        CREATE TABLE IF NOT EXISTS contacts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER,
            contact_type TEXT,  -- 'legal'法人, 'secretary'书记, 'daily'日常联系人
            name TEXT,
            position TEXT,
            phone TEXT,
            is_primary BOOLEAN DEFAULT 0,
            FOREIGN KEY (company_id) REFERENCES companies (id)
        )
    ''')
    
    # 项目信息表
    conn.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER,
            project_name TEXT,
            project_description TEXT,
            total_investment REAL DEFAULT 0,
            design_capacity REAL DEFAULT 0,
            expected_capacity REAL DEFAULT 0,
            actual_capacity REAL DEFAULT 0,
            expected_output REAL DEFAULT 0,
            actual_output REAL DEFAULT 0,
            project_status TEXT,  -- '在建', '投产', '规划中'
            start_date TEXT,
            production_date TEXT,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (company_id) REFERENCES companies (id)
        )
    ''')
    
    # 项目进度更新记录表
    conn.execute('''
        CREATE TABLE IF NOT EXISTS progress_updates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            progress_content TEXT,
            update_date TEXT,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (project_id) REFERENCES projects (id)
        )
    ''')
    
    # 月度经济数据表
    conn.execute('''
        CREATE TABLE IF NOT EXISTS monthly_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER,
            data_type TEXT,  -- 'output'产值, 'capacity'产能, 'tax'税收, 'investment'固定投资, 'added_value'增加值
            year INTEGER,
            month INTEGER,
            planned_value REAL DEFAULT 0,
            actual_value REAL DEFAULT 0,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (company_id) REFERENCES companies (id),
            UNIQUE(company_id, data_type, year, month)
        )
    ''')
    
    # 年度经济数据表
    conn.execute('''
        CREATE TABLE IF NOT EXISTS annual_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_id INTEGER,
            data_type TEXT,  -- 'output'产值, 'capacity'产能, 'tax'税收, 'investment'固定投资, 'added_value'增加值
            year INTEGER,
            planned_value REAL DEFAULT 0,
            actual_value REAL DEFAULT 0,
            created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (company_id) REFERENCES companies (id),
            UNIQUE(company_id, data_type, year)
        )
    ''')
    
    conn.commit()
    conn.close()

# ==================== 页面路由 ====================

@app.route('/')
def index():
    """企业总览看板"""
    conn = get_db_connection()
    companies = conn.execute('''
        SELECT c.*, 
               (SELECT phone FROM contacts WHERE company_id = c.id AND contact_type = 'daily' AND is_primary = 1 LIMIT 1) as contact_phone
        FROM companies c 
        ORDER BY c.created_date DESC
    ''').fetchall()
    conn.close()
    return render_template('index.html', companies=companies, current_year=2025)

@app.route('/company/<int:company_id>')
def company_detail(company_id):
    """企业档案中心 - 显示2025年数据"""
    conn = get_db_connection()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    
    # 获取2025年度汇总数据
    annual_data = conn.execute('''
        SELECT data_type, SUM(actual_value) as total_value
        FROM annual_data 
        WHERE company_id = ? AND year = 2025
        GROUP BY data_type
    ''', (company_id,)).fetchall()
    
    # 获取项目信息
    project = conn.execute('SELECT * FROM projects WHERE company_id = ? LIMIT 1', (company_id,)).fetchone()
    
    conn.close()
    
    annual_dict = {row['data_type']: row['total_value'] or 0 for row in annual_data}
    
    return render_template('company_detail.html', 
                         company=company, 
                         annual_data=annual_dict,
                         project=project,
                         current_year=2025)

@app.route('/contacts')
def contacts_page():
    """通讯录页面"""
    contact_type = request.args.get('type', 'legal')
    conn = get_db_connection()
    
    if contact_type == 'legal':
        contacts = conn.execute('''
            SELECT con.id, c.name as company_name, con.name, con.phone, con.position
            FROM contacts con
            JOIN companies c ON con.company_id = c.id
            WHERE con.contact_type = 'legal'
        ''').fetchall()
    elif contact_type == 'secretary':
        contacts = conn.execute('''
            SELECT con.id, c.name as company_name, con.name, con.phone, con.position
            FROM contacts con
            JOIN companies c ON con.company_id = c.id
            WHERE con.contact_type = 'secretary'
        ''').fetchall()
    else:  # daily
        contacts = conn.execute('''
            SELECT con.id, c.name as company_name, con.name, con.phone, con.position
            FROM contacts con
            JOIN companies c ON con.company_id = c.id
            WHERE con.contact_type = 'daily' AND con.is_primary = 1
        ''').fetchall()
    
    conn.close()
    return render_template('contacts.html', contacts=contacts, current_type=contact_type, current_year=2025)

@app.route('/company/<int:company_id>/economic')
def economic_data(company_id):
    """经济数据页面 - 2025年默认显示，完整月度数据"""
    data_type = request.args.get('type', 'output')
    year = request.args.get('year', type=int, default=2025)  # 默认显示2025年
    
    conn = get_db_connection()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    
    # 获取所有数据类型的数据
    all_monthly_data = {}
    data_types = ['output', 'capacity', 'tax', 'investment', 'added_value']
    
    for dt in data_types:
        monthly_data = conn.execute('''
            SELECT month, planned_value, actual_value
            FROM monthly_data 
            WHERE company_id = ? AND data_type = ? AND year = ?
            ORDER BY month
        ''', (company_id, dt, year)).fetchall()
        all_monthly_data[dt] = monthly_data
    
    # 计算累计数据
    cumulative_data = {}
    for dt in data_types:
        cumulative_planned = 0
        cumulative_actual = 0
        dt_cumulative = []
        
        for i, month_data in enumerate(all_monthly_data[dt]):
            cumulative_planned += month_data['planned_value'] or 0
            cumulative_actual += month_data['actual_value'] or 0
            
            dt_cumulative.append({
                'month': month_data['month'],
                'cumulative_planned': cumulative_planned,
                'cumulative_actual': cumulative_actual
            })
        
        cumulative_data[dt] = dt_cumulative
    
    # 计算季度数据
    quarterly_data = {}
    for dt in data_types:
        dt_quarterly = []
        for quarter in range(1, 5):
            start_month = (quarter - 1) * 3 + 1
            end_month = quarter * 3
            
            quarter_data = conn.execute('''
                SELECT SUM(planned_value) as planned_sum, SUM(actual_value) as actual_sum
                FROM monthly_data 
                WHERE company_id = ? AND data_type = ? AND year = ? 
                AND month BETWEEN ? AND ?
            ''', (company_id, dt, year, start_month, end_month)).fetchone()
            
            dt_quarterly.append({
                'quarter': quarter,
                'planned_sum': quarter_data['planned_sum'] or 0,
                'actual_sum': quarter_data['actual_sum'] or 0
            })
        quarterly_data[dt] = dt_quarterly
    
    conn.close()
    
    # 准备月度数据表格
    months_chinese = ['1月', '2月', '3月', '4月', '5月', '6月', 
                     '7月', '8月', '9月', '10月', '11月', '12月']
    
    data_table = []
    for i in range(12):
        month_num = i + 1
        
        # 为每种数据类型获取数据
        month_data = {}
        for dt in data_types:
            dt_data = next((m for m in all_monthly_data[dt] if m['month'] == month_num), None)
            dt_cumulative = next((c for c in cumulative_data[dt] if c['month'] == month_num), {})
            
            month_data[f'{dt}_planned'] = dt_data['planned_value'] if dt_data else 0
            month_data[f'{dt}_actual'] = dt_data['actual_value'] if dt_data else 0
            month_data[f'{dt}_cumulative_planned'] = dt_cumulative.get('cumulative_planned', 0)
            month_data[f'{dt}_cumulative_actual'] = dt_cumulative.get('cumulative_actual', 0)
        
        data_table.append({
            'month': months_chinese[i],
            'month_num': month_num,
            **month_data
        })
    
    # 生成图表数据
    chart_data = {
        'months': months_chinese,
        'planned': [data[f'{data_type}_planned'] for data in data_table],
        'actual': [data[f'{data_type}_actual'] for data in data_table],
        'cumulative_actual': [data[f'{data_type}_cumulative_actual'] for data in data_table]
    }
    
    return render_template('economic_data.html', 
                         company=company,
                         data=data_table,
                         chart_data=chart_data,
                         data_type=data_type,
                         year=year,
                         company_id=company_id,
                         quarterly_data=quarterly_data,
                         data_types=data_types,
                         years=[2023, 2024, 2025, 2026],
                         current_year=2025,
                         now=datetime.now())

@app.route('/company/<int:company_id>/project')
def project_info(company_id):
    """项目简介页面"""
    conn = get_db_connection()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    project = conn.execute('SELECT * FROM projects WHERE company_id = ?', (company_id,)).fetchone()
    
    progress_updates = []
    if project:
        progress_updates = conn.execute('''
            SELECT * FROM progress_updates 
            WHERE project_id = ? 
            ORDER BY update_date DESC
        ''', (project['id'],)).fetchall()
    
    conn.close()
    return render_template('project_info.html', 
                         company=company, 
                         project=project,
                         progress_updates=progress_updates,
                         current_year=2025)

@app.route('/export')
def export_center():
    """数据调度导出中心"""
    conn = get_db_connection()
    companies = conn.execute('SELECT id, name FROM companies ORDER BY name').fetchall()
    conn.close()
    return render_template('export_center.html', companies=companies, now=datetime.now(), current_year=2025)

@app.route('/company/<int:company_id>/annual_comparison')
def annual_comparison(company_id):
    """年度数据对比页面"""
    conn = get_db_connection()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    
    # 获取近4年的年度数据
    years = [2023, 2024, 2025, 2026]
    comparison_data = {}
    
    for data_type in ['output', 'capacity', 'tax', 'investment', 'added_value']:
        data_by_year = {}
        for year in years:
            annual_data = conn.execute('''
                SELECT planned_value, actual_value 
                FROM annual_data 
                WHERE company_id = ? AND data_type = ? AND year = ?
            ''', (company_id, data_type, year)).fetchone()
            
            if annual_data:
                data_by_year[year] = {
                    'planned': annual_data['planned_value'] or 0,
                    'actual': annual_data['actual_value'] or 0
                }
            else:
                data_by_year[year] = {'planned': 0, 'actual': 0}
        
        comparison_data[data_type] = data_by_year
    
    conn.close()
    
    return render_template('annual_comparison.html', 
                         company=company, 
                         comparison_data=comparison_data,
                         years=years,
                         current_year=2025)

@app.route('/company/<int:company_id>/comprehensive')
def enterprise_comprehensive(company_id):
    """企业综合数据页面 - 修复同比增长计算"""
    # 获取月份参数，默认为当前月份
    selected_month = request.args.get('month', type=int, default=datetime.now().month)
    
    conn = get_db_connection()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    
    # 获取所有数据类型的数据（同时包含当月和累计）
    comprehensive_data = {}
    data_types = ['output', 'capacity', 'tax', 'investment', 'added_value']
    
    for dt in data_types:
        # 获取当月数据
        monthly_data = conn.execute('''
            SELECT actual_value as current_month_value
            FROM monthly_data 
            WHERE company_id = ? AND data_type = ? AND year = ? AND month = ?
        ''', (company_id, dt, 2025, selected_month)).fetchone()
        
        # 获取去年同期当月数据
        last_year_monthly = conn.execute('''
            SELECT actual_value as last_year_month_value
            FROM monthly_data 
            WHERE company_id = ? AND data_type = ? AND year = ? AND month = ?
        ''', (company_id, dt, 2024, selected_month)).fetchone()
        
        # 获取累计数据（1月到当前月份）
        cumulative_data = conn.execute('''
            SELECT SUM(actual_value) as cumulative_value
            FROM monthly_data 
            WHERE company_id = ? AND data_type = ? AND year = ? AND month <= ?
        ''', (company_id, dt, 2025, selected_month)).fetchone()
        
        # 获取去年同期累计数据
        last_year_cumulative = conn.execute('''
            SELECT SUM(actual_value) as last_year_cumulative_value
            FROM monthly_data 
            WHERE company_id = ? AND data_type = ? AND year = ? AND month <= ?
        ''', (company_id, dt, 2024, selected_month)).fetchone()
        
        # 当月数据
        current_month_value = monthly_data['current_month_value'] if monthly_data and monthly_data['current_month_value'] is not None else 0
        last_year_month_value = last_year_monthly['last_year_month_value'] if last_year_monthly and last_year_monthly['last_year_month_value'] is not None else 0
        
        # 修复同比增长率计算
        month_growth_rate = 0
        if last_year_month_value > 0:
            month_growth_rate = ((current_month_value - last_year_month_value) / last_year_month_value) * 100
        elif last_year_month_value == 0 and current_month_value > 0:
            month_growth_rate = 100  # 去年为0，今年有值，增长100%
        elif last_year_month_value == 0 and current_month_value == 0:
            month_growth_rate = 0    # 两年都为0，增长0%
        else:
            month_growth_rate = -100 if current_month_value == 0 else ((current_month_value - last_year_month_value) / abs(last_year_month_value)) * 100
        
        # 累计数据
        cumulative_value = cumulative_data['cumulative_value'] if cumulative_data and cumulative_data['cumulative_value'] is not None else 0
        last_year_cumulative_value = last_year_cumulative['last_year_cumulative_value'] if last_year_cumulative and last_year_cumulative['last_year_cumulative_value'] is not None else 0
        
        # 修复累计同比增长率计算
        cumulative_growth_rate = 0
        if last_year_cumulative_value > 0:
            cumulative_growth_rate = ((cumulative_value - last_year_cumulative_value) / last_year_cumulative_value) * 100
        elif last_year_cumulative_value == 0 and cumulative_value > 0:
            cumulative_growth_rate = 100  # 去年为0，今年有值，增长100%
        elif last_year_cumulative_value == 0 and cumulative_value == 0:
            cumulative_growth_rate = 0    # 两年都为0，增长0%
        else:
            cumulative_growth_rate = -100 if cumulative_value == 0 else ((cumulative_value - last_year_cumulative_value) / abs(last_year_cumulative_value)) * 100
        
        comprehensive_data[dt] = {
            'current_month_value': current_month_value,
            'cumulative_value': cumulative_value,
            'month_growth_rate': month_growth_rate,
            'cumulative_growth_rate': cumulative_growth_rate
        }
        # 添加调试信息
    print(f"=== 综合数据调试信息 ===")
    print(f"企业ID: {company_id}, 月份: {selected_month}")
    for dt in data_types:
        print(f"{dt}: {comprehensive_data[dt]}")
    print("=======================")
    conn.close()
    
    # 月份中文名称
    months_chinese = ['', '1月', '2月', '3月', '4月', '5月', '6月', 
                     '7月', '8月', '9月', '10月', '11月', '12月']
    
    return render_template('enterprise_comprehensive.html',
                         company=company,
                         comprehensive_data=comprehensive_data,
                         selected_month=selected_month,
                         current_month_name=months_chinese[selected_month],
                         months=range(1, 13),
                         months_chinese=months_chinese,
                         current_year=2025)

# ==================== API 接口 ====================

@app.route('/api/company', methods=['POST'])
def add_company():
    """添加企业"""
    data = request.get_json()
    conn = get_db_connection()
    try:
        cursor = conn.execute('''
            INSERT INTO companies (name, legal_person, main_products, product_model, 
                                 party_secretary, total_investment, employee_count, 
                                 register_date, completion_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            data['name'], data['legal_person'], data['main_products'],
            data.get('product_model', ''), data.get('party_secretary', ''),
            data.get('total_investment', 0), data.get('employee_count', 0),
            data.get('register_date', ''), data.get('completion_date', '')
        ))
        
        company_id = cursor.lastrowid
        
        # 添加联系人
        if 'contacts' in data:
            for contact in data['contacts']:
                conn.execute('''
                    INSERT INTO contacts (company_id, contact_type, name, position, phone, is_primary)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (company_id, contact['type'], contact['name'], 
                      contact.get('position', ''), contact.get('phone', ''), 
                      contact.get('is_primary', 0)))
        
        conn.commit()
        return jsonify({'message': '企业添加成功!', 'success': True, 'company_id': company_id})
    except Exception as e:
        return jsonify({'message': f'添加失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/company/<int:company_id>', methods=['DELETE'])
def delete_company(company_id):
    """删除企业"""
    conn = get_db_connection()
    try:
        conn.execute('DELETE FROM companies WHERE id = ?', (company_id,))
        conn.commit()
        return jsonify({'message': '企业删除成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'删除失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/economic_data', methods=['POST'])
def update_economic_data():
    """更新经济数据 - 修复乱码问题"""
    conn = get_db_connection()
    try:
        # 确保正确解析JSON数据
        data = request.get_json(force=True)
        print("接收到的经济数据:", data)
        
        # 更新月度数据
        if 'monthly' in data:
            for monthly in data['monthly']:
                conn.execute('''
                    INSERT OR REPLACE INTO monthly_data 
                    (company_id, data_type, year, month, planned_value, actual_value)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    monthly['company_id'], 
                    monthly['data_type'], 
                    monthly['year'],
                    monthly['month'], 
                    monthly['planned_value'], 
                    monthly['actual_value']
                ))
        
        # 更新年度数据
        if 'annual' in data:
            for annual in data['annual']:
                conn.execute('''
                    INSERT OR REPLACE INTO annual_data 
                    (company_id, data_type, year, planned_value, actual_value)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    annual['company_id'], 
                    annual['data_type'], 
                    annual['year'],
                    annual['planned_value'], 
                    annual['actual_value']
                ))
        
        conn.commit()
        return jsonify({'message': '经济数据更新成功!', 'success': True})
    except Exception as e:
        print(f"经济数据更新错误: {str(e)}")
        return jsonify({'message': f'更新失败: {str(e)}', 'success': False})
    finally:
        conn.close()

# ==================== 新增的通讯录API接口 ====================

@app.route('/api/companies', methods=['GET'])
def get_companies():
    """获取所有企业列表"""
    conn = get_db_connection()
    try:
        companies = conn.execute('SELECT id, name FROM companies ORDER BY name').fetchall()
        companies_list = [{'id': company['id'], 'name': company['name']} for company in companies]
        return jsonify(companies_list)
    except Exception as e:
        return jsonify({'message': f'获取企业列表失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/contact', methods=['POST'])
def add_contact():
    """添加联系人"""
    data = request.get_json()
    conn = get_db_connection()
    try:
        conn.execute('''
            INSERT INTO contacts (company_id, contact_type, name, position, phone, is_primary)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (
            data['company_id'], data['contact_type'], data['name'],
            data.get('position', ''), data.get('phone', ''), data.get('is_primary', 0)
        ))
        
        conn.commit()
        return jsonify({'message': '联系人添加成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'添加失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/contact/<int:contact_id>', methods=['PUT'])
def update_contact(contact_id):
    """更新联系人"""
    data = request.get_json()
    conn = get_db_connection()
    try:
        conn.execute('''
            UPDATE contacts SET name = ?, position = ?, phone = ?
            WHERE id = ?
        ''', (data['name'], data.get('position', ''), data.get('phone', ''), contact_id))
        
        conn.commit()
        return jsonify({'message': '联系人更新成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'更新失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/contact/<int:contact_id>', methods=['DELETE'])
def delete_contact(contact_id):
    """删除联系人"""
    conn = get_db_connection()
    try:
        conn.execute('DELETE FROM contacts WHERE id = ?', (contact_id,))
        conn.commit()
        return jsonify({'message': '联系人删除成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'删除失败: {str(e)}', 'success': False})
    finally:
        conn.close()

# ==================== 原有的项目API接口 ====================
@app.route('/api/project', methods=['POST'])
def update_project():
    """更新项目信息 - 完全修复版"""
    data = request.get_json()
    print("接收到的项目数据:", data)
    conn = get_db_connection()
    try:
        # 确保所有数值字段都有值
        company_id = data.get('company_id')
        if not company_id:
            return jsonify({'message': '缺少公司ID', 'success': False})
        
        # 转换所有数值字段
        project_data = {
            'project_name': data.get('project_name', ''),
            'project_description': data.get('project_description', ''),
            'total_investment': float(data.get('total_investment', 0)),
            'design_capacity': float(data.get('design_capacity', 0)),
            'expected_capacity': float(data.get('expected_capacity', 0)),
            'actual_capacity': float(data.get('actual_capacity', 0)),
            'expected_output': float(data.get('expected_output', 0)),
            'actual_output': float(data.get('actual_output', 0)),
            'project_status': data.get('project_status', '在建'),
            'start_date': data.get('start_date', ''),
            'production_date': data.get('production_date', '')
        }
        
        # 检查是否已存在项目
        existing_project = conn.execute('SELECT id FROM projects WHERE company_id = ?', (company_id,)).fetchone()
        
        if existing_project:
            # 更新现有项目
            conn.execute('''
                UPDATE projects SET 
                project_name = ?, project_description = ?, total_investment = ?,
                design_capacity = ?, expected_capacity = ?, actual_capacity = ?,
                expected_output = ?, actual_output = ?, project_status = ?,
                start_date = ?, production_date = ?
                WHERE company_id = ?
            ''', (
                project_data['project_name'], project_data['project_description'],
                project_data['total_investment'], project_data['design_capacity'],
                project_data['expected_capacity'], project_data['actual_capacity'],
                project_data['expected_output'], project_data['actual_output'],
                project_data['project_status'], project_data['start_date'],
                project_data['production_date'], company_id
            ))
        else:
            # 创建新项目
            conn.execute('''
                INSERT INTO projects 
                (company_id, project_name, project_description, total_investment,
                 design_capacity, expected_capacity, actual_capacity,
                 expected_output, actual_output, project_status, start_date, production_date)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                company_id, project_data['project_name'], project_data['project_description'],
                project_data['total_investment'], project_data['design_capacity'],
                project_data['expected_capacity'], project_data['actual_capacity'],
                project_data['expected_output'], project_data['actual_output'],
                project_data['project_status'], project_data['start_date'],
                project_data['production_date']
            ))
        
        conn.commit()
        return jsonify({'message': '项目信息保存成功!', 'success': True})
    except Exception as e:
        print(f"项目保存错误: {str(e)}")
        return jsonify({'message': f'保存失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/progress', methods=['POST'])
def add_progress():
    """添加项目进度"""
    data = request.get_json()
    conn = get_db_connection()
    try:
        conn.execute('''
            INSERT INTO progress_updates (project_id, progress_content, update_date)
            VALUES (?, ?, ?)
        ''', (data['project_id'], data['progress_content'], data['update_date']))
        
        conn.commit()
        return jsonify({'message': '进度更新成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'保存失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/progress/<int:progress_id>', methods=['DELETE'])
def delete_progress(progress_id):
    """删除项目进度 - 新增功能"""
    conn = get_db_connection()
    try:
        conn.execute('DELETE FROM progress_updates WHERE id = ?', (progress_id,))
        conn.commit()
        return jsonify({'message': '进度删除成功!', 'success': True})
    except Exception as e:
        return jsonify({'message': f'删除失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/export_contacts', methods=['POST'])
def export_contacts():
    """导出通讯录 - 修复版"""
    data = request.get_json()
    contact_type = data.get('type', 'legal')
    
    conn = get_db_connection()
    try:
        if contact_type == 'legal':
            contacts = conn.execute('''
                SELECT c.name as company_name, con.name, con.phone, con.position
                FROM contacts con
                JOIN companies c ON con.company_id = c.id
                WHERE con.contact_type = 'legal'
            ''').fetchall()
        elif contact_type == 'secretary':
            contacts = conn.execute('''
                SELECT c.name as company_name, con.name, con.phone, con.position
                FROM contacts con
                JOIN companies c ON con.company_id = c.id
                WHERE con.contact_type = 'secretary'
            ''').fetchall()
        else:  # daily
            contacts = conn.execute('''
                SELECT c.name as company_name, con.name, con.phone, con.position
                FROM contacts con
                JOIN companies c ON con.company_id = c.id
                WHERE con.contact_type = 'daily' AND con.is_primary = 1
            ''').fetchall()
        
        # 创建Excel文件
        output = io.BytesIO()
        
        contact_data = []
        for contact in contacts:
            contact_data.append({
                '企业名称': contact['company_name'],
                '姓名': contact['name'],
                '职位': contact['position'],
                '电话': contact['phone']
            })
        
        df = pd.DataFrame(contact_data)
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='通讯录', index=False)
        
        output.seek(0)
        
        type_name = {'legal': '法人', 'secretary': '书记', 'daily': '日常联系人'}.get(contact_type, '通讯录')
        filename = f'{type_name}通讯录_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        print(f"通讯录导出错误: {str(e)}")
        return jsonify({'message': f'导出失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/export_custom_fields', methods=['POST'])
def export_custom_fields():
    """自定义字段导出 - 完整修复版"""
    data = request.get_json()
    conn = get_db_connection()
    
    try:
        company_ids = data.get('company_ids', [])
        selected_fields = data.get('selected_fields', [])
        export_format = data.get('format', 'excel')
        
        print(f"=== 自定义字段导出调试 ===")
        print(f"企业ID: {company_ids}")
        print(f"选中字段: {selected_fields}")
        print(f"导出格式: {export_format}")
        
        if not company_ids:
            return jsonify({'message': '未选择任何企业', 'success': False})
        
        if not selected_fields:
            return jsonify({'message': '未选择任何字段', 'success': False})
        
        # 收集所有数据
        all_data = []
        
        for company_id in company_ids:
            # 获取企业基本信息
            company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
            if not company:
                continue
                
            row_data = {}
            
            # 基础信息字段
            if 'company_name' in selected_fields:
                row_data['企业名称'] = company['name']
            if 'legal_person' in selected_fields:
                row_data['法定代表人'] = company['legal_person'] or ''
            if 'main_products' in selected_fields:
                row_data['主营产品'] = company['main_products'] or ''
            if 'product_model' in selected_fields:
                row_data['产品型号'] = company['product_model'] or ''
            if 'party_secretary' in selected_fields:
                row_data['党组织书记'] = company['party_secretary'] or ''
            if 'total_investment' in selected_fields:
                row_data['总投资数(万元)'] = company['total_investment'] or 0
            if 'employee_count' in selected_fields:
                row_data['员工人数'] = company['employee_count'] or 0
            if 'register_date' in selected_fields:
                row_data['注册时间'] = company['register_date'] or ''
            if 'completion_date' in selected_fields:
                row_data['建成时间'] = company['completion_date'] or ''
            
            # 联系人信息
            if 'legal_contact' in selected_fields:
                legal_contact = conn.execute('''
                    SELECT name, phone, position FROM contacts 
                    WHERE company_id = ? AND contact_type = 'legal' LIMIT 1
                ''', (company_id,)).fetchone()
                if legal_contact:
                    row_data['法人代表'] = f"{legal_contact['name']} ({legal_contact['position']})"
                    row_data['法人电话'] = legal_contact['phone']
            
            if 'secretary_contact' in selected_fields:
                secretary_contact = conn.execute('''
                    SELECT name, phone, position FROM contacts 
                    WHERE company_id = ? AND contact_type = 'secretary' LIMIT 1
                ''', (company_id,)).fetchone()
                if secretary_contact:
                    row_data['党组织书记'] = f"{secretary_contact['name']} ({secretary_contact['position']})"
                    row_data['书记电话'] = secretary_contact['phone']
            
            if 'daily_contact' in selected_fields:
                daily_contact = conn.execute('''
                    SELECT name, phone, position FROM contacts 
                    WHERE company_id = ? AND contact_type = 'daily' AND is_primary = 1 LIMIT 1
                ''', (company_id,)).fetchone()
                if daily_contact:
                    row_data['日常联系人'] = f"{daily_contact['name']} ({daily_contact['position']})"
                    row_data['联系人电话'] = daily_contact['phone']
            
            # 项目信息
            if any(field in selected_fields for field in ['project_name', 'project_description', 'project_status']):
                project = conn.execute('SELECT * FROM projects WHERE company_id = ?', (company_id,)).fetchone()
                if project:
                    if 'project_name' in selected_fields:
                        row_data['项目名称'] = project['project_name'] or ''
                    if 'project_description' in selected_fields:
                        row_data['项目简介'] = project['project_description'] or ''
                    if 'project_status' in selected_fields:
                        row_data['项目状态'] = project['project_status'] or ''
                    if 'total_investment' in selected_fields:
                        row_data['项目投资(万元)'] = project['total_investment'] or 0
                    if 'start_date' in selected_fields:
                        row_data['开工时间'] = project['start_date'] or ''
                    if 'production_date' in selected_fields:
                        row_data['投产时间'] = project['production_date'] or ''
            
            # 经济数据 - 2025年度数据
            economic_fields = ['output', 'capacity', 'tax', 'investment', 'added_value']
            for field in economic_fields:
                if field in selected_fields:
                    annual_data = conn.execute('''
                        SELECT planned_value, actual_value FROM annual_data 
                        WHERE company_id = ? AND data_type = ? AND year = 2025
                    ''', (company_id, field)).fetchone()
                    
                    field_names = {
                        'output': '产值',
                        'capacity': '产能',
                        'tax': '税收',
                        'investment': '投资',
                        'added_value': '增加值'
                    }
                    
                    if annual_data:
                        row_data[f'2025{field_names[field]}计划(万元)'] = annual_data['planned_value'] or 0
                        row_data[f'2025{field_names[field]}实际(万元)'] = annual_data['actual_value'] or 0
            
            all_data.append(row_data)
        
        print(f"准备导出的数据条数: {len(all_data)}")
        
        # 根据格式导出
        if export_format == 'excel':
            return export_to_excel(all_data, '企业自定义数据')
        elif export_format == 'word':
            return export_to_word(all_data, '企业自定义数据')
        elif export_format == 'pdf':
            return export_to_pdf(all_data, '企业自定义数据')
        else:
            return jsonify({'message': '不支持的导出格式', 'success': False})
    
    except Exception as e:
        print(f"自定义字段导出错误: {str(e)}")
        return jsonify({'message': f'导出失败: {str(e)}', 'success': False})
    finally:
        conn.close()

@app.route('/api/export_advanced', methods=['POST'])
def export_advanced():
    """高级导出功能 - 支持Excel、Word、PDF三种格式"""
    data = request.get_json()
    conn = get_db_connection()
    
    try:
        company_ids = data.get('company_ids', [])
        export_format = data.get('format', 'excel')
        
        # 获取选中的字段
        basic_fields = data.get('basic_fields', [])
        time_fields = data.get('time_fields', [])
        project_fields = data.get('project_fields', [])
        economic_data = data.get('economic_data', {})
        
        print(f"=== 高级导出调试信息 ===")
        print(f"接收到的 company_ids: {company_ids}")
        print(f"导出格式: {export_format}")
        print(f"基础字段: {basic_fields}")
        print(f"时间字段: {time_fields}")
        print(f"项目字段: {project_fields}")
        print(f"经济数据: {economic_data}")
        
        if not company_ids:
            return jsonify({'error': '未选择任何企业'}), 400
        
        # 修复查询逻辑
        query = 'SELECT * FROM companies'
        params = []
        
        if company_ids and isinstance(company_ids, list) and len(company_ids) > 0:
            company_ids = [int(cid) for cid in company_ids if cid is not None and str(cid).isdigit()]
            if company_ids:
                placeholders = ','.join('?' * len(company_ids))
                query += f' WHERE id IN ({placeholders})'
                params = company_ids
        
        companies = conn.execute(query, params).fetchall()
        print(f"实际导出的企业数量: {len(companies)}")
        
        # 收集所有数据
        all_data = []
        
        for company in companies:
            row_data = {}
            
            # 基础信息
            if 'company_name' in basic_fields:
                row_data['企业名称'] = company['name']
            if 'legal_person' in basic_fields:
                row_data['法定代表人'] = company['legal_person'] or ''
            if 'main_products' in basic_fields:
                row_data['主营产品'] = company['main_products'] or ''
            if 'product_model' in basic_fields:
                row_data['产品型号'] = company['product_model'] or ''
            if 'party_secretary' in basic_fields:
                row_data['党组织书记'] = company['party_secretary'] or ''
            if 'total_investment' in basic_fields:
                row_data['总投资数(万元)'] = company['total_investment'] or 0
            if 'employee_count' in basic_fields:
                row_data['员工人数'] = company['employee_count'] or 0
            
            # 时间信息
            if 'register_date' in time_fields:
                row_data['注册时间'] = company['register_date'] or ''
            if 'completion_date' in time_fields:
                row_data['建成时间'] = company['completion_date'] or ''
            
            # 项目信息
            project = conn.execute('SELECT * FROM projects WHERE company_id = ?', (company['id'],)).fetchone()
            if project:
                if 'project_description' in project_fields:
                    row_data['项目简介'] = project['project_description'] or ''
                if 'start_date' in project_fields:
                    row_data['开工时间'] = project['start_date'] or ''
                if 'production_date' in project_fields:
                    row_data['投产时间'] = project['production_date'] or ''
            
            # 经济数据
            if economic_data:
                data_type_map = {
                    'output': '产值',
                    'tax': '税收', 
                    'investment': '投资',
                    'added_value': '增加值',
                    'capacity': '产能'
                }
                
                for data_type, settings in economic_data.items():
                    if settings.get('selected'):
                        time_type = settings.get('time_type', 'monthly')
                        year = settings.get('year', 2025)
                        month = settings.get('month')
                        quarter = settings.get('quarter')
                        
                        if time_type == 'monthly' and month:
                            # 月度数据
                            monthly_data = conn.execute('''
                                SELECT planned_value, actual_value
                                FROM monthly_data 
                                WHERE company_id = ? AND data_type = ? AND year = ? AND month = ?
                            ''', (company['id'], data_type, year, month)).fetchone()
                            
                            if monthly_data:
                                row_data[f'{data_type_map[data_type]}计划值(万元)'] = monthly_data['planned_value'] or 0
                                row_data[f'{data_type_map[data_type]}实际值(万元)'] = monthly_data['actual_value'] or 0
                        
                        elif time_type == 'quarterly' and quarter:
                            # 季度数据
                            start_month = (quarter - 1) * 3 + 1
                            end_month = quarter * 3
                            
                            quarterly_data = conn.execute('''
                                SELECT SUM(planned_value) as planned_sum, SUM(actual_value) as actual_sum
                                FROM monthly_data 
                                WHERE company_id = ? AND data_type = ? AND year = ? 
                                AND month BETWEEN ? AND ?
                            ''', (company['id'], data_type, year, start_month, end_month)).fetchone()
                            
                            if quarterly_data:
                                row_data[f'Q{quarter}{data_type_map[data_type]}计划值(万元)'] = quarterly_data['planned_sum'] or 0
                                row_data[f'Q{quarter}{data_type_map[data_type]}实际值(万元)'] = quarterly_data['actual_sum'] or 0
                        
                        elif time_type == 'annual':
                            # 年度数据
                            annual_data = conn.execute('''
                                SELECT planned_value, actual_value
                                FROM annual_data 
                                WHERE company_id = ? AND data_type = ? AND year = ?
                            ''', (company['id'], data_type, year)).fetchone()
                            
                            if annual_data:
                                row_data[f'{year}年{data_type_map[data_type]}计划值(万元)'] = annual_data['planned_value'] or 0
                                row_data[f'{year}年{data_type_map[data_type]}实际值(万元)'] = annual_data['actual_value'] or 0
            
            all_data.append(row_data)
        
        # 根据格式导出
        if export_format == 'excel':
            return export_to_excel(all_data, '企业数据导出')
        elif export_format == 'word':
            return export_to_word(all_data, '企业数据导出')
        elif export_format == 'pdf':
            return export_to_pdf(all_data, '企业数据导出')
        else:
            return jsonify({'message': '不支持的导出格式', 'success': False})
    
    except Exception as e:
        print(f"高级导出错误: {str(e)}")
        return jsonify({'message': f'导出失败: {str(e)}', 'success': False})
    finally:
        conn.close()

def export_to_excel(data, title):
    """导出为Excel格式 - 修复中文乱码"""
    output = io.BytesIO()
    
    if data:
        # 创建DataFrame时指定编码
        df = pd.DataFrame(data)
        
        # 使用openpyxl引擎并设置中文支持
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=title, index=False)
            
            # 获取工作表并设置列宽
            worksheet = writer.sheets[title]
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                worksheet.column_dimensions[column_letter].width = adjusted_width
    else:
        # 创建空的工作表
        df = pd.DataFrame({'提示': ['没有数据可导出']})
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=title, index=False)
    
    output.seek(0)
    filename = f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

def export_to_word(data, title):
    """导出为Word格式 - 简化修复版"""
    document = Document()
    
    # 添加标题
    document.add_heading(title, 0)
    document.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    document.add_paragraph()
    
    if data:
        # 添加表格
        table = document.add_table(rows=1, cols=len(data[0].keys()))
        table.style = 'Table Grid'
        
        # 添加表头
        headers = list(data[0].keys())
        for i, header in enumerate(headers):
            table.cell(0, i).text = str(header)
        
        # 添加数据行
        for row_data in data:
            row_cells = table.add_row().cells
            for i, key in enumerate(headers):
                row_cells[i].text = str(row_data.get(key, ''))
    else:
        document.add_paragraph('没有数据可导出')
    
    # 保存到内存
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    
    filename = f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

def export_to_pdf(data, title):
    """导出为PDF格式 - 简化修复版"""
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []
    
    # 定义样式
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # 居中
    )
    
    # 添加标题
    elements.append(Paragraph(title, title_style))
    elements.append(Paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    if data:
        # 准备表格数据
        table_data = []
        
        # 表头
        headers = list(data[0].keys())
        table_data.append(headers)
        
        # 数据行
        for row_data in data:
            row = [str(row_data.get(key, '')) for key in headers]
            table_data.append(row)
        
        # 创建表格
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
    else:
        elements.append(Paragraph('没有数据可导出', styles['Normal']))
    
    # 构建PDF
    doc.build(elements)
    output.seek(0)
    
    filename = f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    
    return send_file(
        output,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )

# 添加示例数据
def add_sample_data():
    """添加示例数据用于测试"""
    conn = get_db_connection()
    
    # 检查是否已有数据
    existing = conn.execute('SELECT COUNT(*) as count FROM companies').fetchone()
    if existing['count'] > 0:
        conn.close()
        return
    
    # 添加示例企业
    companies_data = [
        ('三一能源装备有限公司', '张三', '风电设备', 'SE-3000', '李四', 15000.0, 200, '2020-01-15', '2022-06-30'),
        ('天能重工集团有限公司', '王五', '光伏组件', 'TN-500W', '赵六', 8000.0, 150, '2019-05-20', '2021-12-15'),
        ('华电新能源有限公司', '陈七', '储能设备', 'HD-100KWH', '钱八', 12000.0, 180, '2021-03-10', '2023-08-20')
    ]
    
    for company in companies_data:
        cursor = conn.execute('''
            INSERT INTO companies (name, legal_person, main_products, product_model, 
                                 party_secretary, total_investment, employee_count, 
                                 register_date, completion_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', company)
        
        company_id = cursor.lastrowid
        
        # 添加联系人
        conn.execute('''
            INSERT INTO contacts (company_id, contact_type, name, position, phone, is_primary)
            VALUES (?, 'legal', ?, '法定代表人', '13800138000', 1)
        ''', (company_id, company[1]))
        
        conn.execute('''
            INSERT INTO contacts (company_id, contact_type, name, position, phone, is_primary)
            VALUES (?, 'secretary', ?, '党组织书记', '13900139000', 1)
        ''', (company_id, company[4]))
        
        conn.execute('''
            INSERT INTO contacts (company_id, contact_type, name, position, phone, is_primary)
            VALUES (?, 'daily', '周经理', '业务经理', '13700137000', 1)
        ''', (company_id,))
        
        # 添加项目信息
        conn.execute('''
            INSERT INTO projects (company_id, project_name, project_description, total_investment,
                               design_capacity, expected_capacity, actual_capacity,
                               project_status, start_date, production_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            company_id, 
            f'{company[0]}主要项目',
            f'{company[2]}生产线建设项目',
            company[5] * 0.8,
            company[5] * 0.1,
            company[5] * 0.12,
            company[5] * 0.1,
            '投产',
            '2022-01-01',
            '2023-06-01'
        ))
        
        # 添加示例月度数据
        for data_type in ['output', 'capacity', 'tax', 'investment', 'added_value']:
            for year in [2023, 2024, 2025, 2026]:
                # 年度数据
                base_value = company[5]
                if data_type == 'output':
                    annual_planned = base_value * 0.8
                    annual_actual = base_value * 0.9
                elif data_type == 'capacity':
                    annual_planned = base_value * 0.1
                    annual_actual = base_value * 0.12
                elif data_type == 'tax':
                    annual_planned = base_value * 0.05
                    annual_actual = base_value * 0.06
                elif data_type == 'investment':
                    annual_planned = base_value * 0.3
                    annual_actual = base_value * 0.35
                else:  # added_value
                    annual_planned = base_value * 0.15
                    annual_actual = base_value * 0.18
                
                conn.execute('''
                    INSERT INTO annual_data (company_id, data_type, year, planned_value, actual_value)
                    VALUES (?, ?, ?, ?, ?)
                ''', (company_id, data_type, year, annual_planned, annual_actual))
                
                # 月度数据
                for month in range(1, 13):
                    if year == 2025:
                        # 2025年数据更详细
                        growth_factor = 1 + (month * 0.02)
                        monthly_planned = annual_planned / 12 * growth_factor
                        monthly_actual = annual_actual / 12 * growth_factor * 0.95
                    else:
                        monthly_planned = annual_planned / 12
                        monthly_actual = annual_actual / 12
                    
                    conn.execute('''
                        INSERT INTO monthly_data (company_id, data_type, year, month, planned_value, actual_value)
                        VALUES (?, ?, ?, ?, ?, ?)
                    ''', (company_id, data_type, year, month, monthly_planned, monthly_actual))
    
    conn.commit()
    conn.close()
    print("示例数据添加完成！")

if __name__ == '__main__':
    init_db()
    add_sample_data()  # 添加示例数据
    print("=" * 60)
    print("能源装备企业一企一档管理系统 V2.0 - 2025年数据版本 启动成功!")
    print("访问地址: http://localhost:5000")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)